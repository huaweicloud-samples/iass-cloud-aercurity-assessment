import os
import json
import shutil
from datetime import datetime
from typing import List, Dict, Any

from docx import Document
from docx.oxml.ns import qn
import pandas as pd

from config import TEMPLATE_DIR, UPLOAD_DIR, ALLOWED_EXTENSIONS


class DocumentHandler:
    """Word文档处理工具类"""

    @staticmethod
    def validate_file_extension(filename: str) -> bool:
        """校验文件扩展名是否合法"""
        if not filename:
            return False
        ext = os.path.splitext(filename)[1].lower()
        return ext in ALLOWED_EXTENSIONS

    @staticmethod
    def parse_word_document(file_path: str) -> List[Dict[str, Any]]:
        """解析Word文档，按文档顺序提取段落和表格，保留格式信息"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        doc = Document(file_path)
        contents = []

        # 按文档顺序遍历所有body子元素
        from docx.oxml.ns import qn as _qn
        body = doc.element.body

        para_idx = 0
        table_idx = 0

        for child in body:
            # 段落
            if child.tag == _qn('w:p'):
                if para_idx < len(doc.paragraphs):
                    para = doc.paragraphs[para_idx]
                    para_idx += 1

                    # 提取格式信息
                    style_name = para.style.name if para.style else "Normal"
                    style_id = para.style.style_id if para.style else "Normal"

                    # 判断标题层级
                    heading_level = 0
                    if style_name.startswith("Heading") or style_name.startswith("heading"):
                        try:
                            heading_level = int(''.join(c for c in style_name if c.isdigit()))
                        except:
                            heading_level = 1
                    elif style_id.startswith("Heading") or style_id.startswith("heading"):
                        try:
                            heading_level = int(''.join(c for c in style_id if c.isdigit()))
                        except:
                            heading_level = 1

                    # 提取run级别的格式
                    runs_info = []
                    for run in para.runs:
                        run_info = {
                            "text": run.text,
                            "bold": run.bold,
                            "italic": run.italic,
                            "underline": run.underline is not None and run.underline is not False,
                            "font_size": run.font.size.pt if run.font.size and run.font.size.pt else None,
                            "font_name": run.font.name,
                        }
                        runs_info.append(run_info)

                    # 判断列表类型
                    is_list = False
                    list_level = 0
                    numPr = child.find(_qn('w:pPr') + '/' + _qn('w:numPr'))
                    if numPr is not None:
                        is_list = True
                        ilvl = numPr.find(_qn('w:ilvl'))
                        if ilvl is not None:
                            try:
                                list_level = int(ilvl.get(_qn('w:val'), '0'))
                            except:
                                list_level = 0

                    # 对齐方式
                    alignment = None
                    pPr = child.find(_qn('w:pPr'))
                    if pPr is not None:
                        jc = pPr.find(_qn('w:jc'))
                        if jc is not None:
                            alignment = jc.get(_qn('w:val'))

                    text = para.text.strip()
                    if text or runs_info:
                        contents.append({
                            "content_type": "paragraph",
                            "content_data": json.dumps({
                                "text": text,
                                "style": style_name,
                                "heading_level": heading_level,
                                "is_list": is_list,
                                "list_level": list_level,
                                "alignment": alignment,
                                "runs": runs_info,
                            }, ensure_ascii=False)
                        })

            # 表格
            elif child.tag == _qn('w:tbl'):
                if table_idx < len(doc.tables):
                    table = doc.tables[table_idx]
                    table_idx += 1

                    table_data = []
                    merge_info = []  # 存储每个单元格的合并信息
                    for ri, row in enumerate(table.rows):
                        row_data = []
                        row_merge = []
                        for ci, cell in enumerate(row.cells):
                            row_data.append(cell.text.strip())
                            # 提取单元格合并信息
                            tc = cell._tc
                            tcPr = tc.find(_qn('w:tcPr'))
                            cell_merge = {}
                            if tcPr is not None:
                                vMerge = tcPr.find(_qn('w:vMerge'))
                                hMerge = tcPr.find(_qn('w:hMerge'))
                                gridSpan = tcPr.find(_qn('w:gridSpan'))
                                if vMerge is not None:
                                    cell_merge['vMerge'] = vMerge.get(_qn('w:val'), 'continue')
                                if hMerge is not None:
                                    cell_merge['hMerge'] = hMerge.get(_qn('w:val'), 'continue')
                                if gridSpan is not None:
                                    cell_merge['gridSpan'] = int(gridSpan.get(_qn('w:val'), '1'))
                            row_merge.append(cell_merge)
                        table_data.append(row_data)
                        merge_info.append(row_merge)

                    # 提取表格样式
                    tbl_style = ""
                    tblPr = child.find(_qn('w:tblPr'))
                    if tblPr is not None:
                        tblStyleEl = tblPr.find(_qn('w:tblStyle'))
                        if tblStyleEl is not None:
                            tbl_style = tblStyleEl.get(_qn('w:val'), "")

                    contents.append({
                        "content_type": "table",
                        "content_data": json.dumps({
                            "rows": len(table.rows),
                            "cols": len(table.columns),
                            "data": table_data,
                            "style": tbl_style,
                            "merge_info": merge_info,  # 添加合并信息
                        }, ensure_ascii=False)
                    })

        return contents

    @staticmethod
    def create_document_from_template(template_path: str, output_path: str) -> str:
        """基于模板创建新文档副本"""
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板文件不存在: {template_path}")

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        shutil.copy2(template_path, output_path)
        return output_path

    @staticmethod
    def update_document_content(file_path: str, content_data: str) -> bool:
        """更新Word文档内容"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        doc = Document(file_path)
        data = json.loads(content_data)

        if "paragraphs" in data:
            for i, para_data in enumerate(data["paragraphs"]):
                if i < len(doc.paragraphs):
                    doc.paragraphs[i].text = para_data.get("text", "")

        doc.save(file_path)
        return True

    @staticmethod
    def get_document_info(file_path: str) -> Dict[str, Any]:
        """获取文档基本信息"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        doc = Document(file_path)
        return {
            "paragraphs_count": len(doc.paragraphs),
            "tables_count": len(doc.tables),
            "sections_count": len(doc.sections),
        }

    @staticmethod
    def generate_word_document_from_contents(
        template_path: str,
        contents: List[Dict[str, Any]],
        output_path: str
    ) -> bool:
        """根据编辑内容生成Word文档，保持模板格式"""
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板文件不存在: {template_path}")

        # 复制模板文件到输出路径
        shutil.copy2(template_path, output_path)
        doc = Document(output_path)

        # 按文档顺序遍历所有body子元素
        from docx.oxml.ns import qn as _qn
        body = doc.element.body

        para_idx = 0
        table_idx = 0
        content_idx = 0

        for child in body:
            if content_idx >= len(contents):
                break

            content = contents[content_idx]

            # 处理段落
            if child.tag == _qn('w:p'):
                if para_idx < len(doc.paragraphs):
                    para = doc.paragraphs[para_idx]
                    para_idx += 1

                    # 解析内容数据
                    try:
                        if isinstance(content["content_data"], str):
                            data = json.loads(content["content_data"])
                        else:
                            data = content["content_data"]

                        # 清空段落内容（保留样式）
                        for run in para.runs[:]:
                            run._element.getparent().remove(run._element)

                        # 添加文本，保留格式
                        text = data.get("text", "")
                        if text:
                            run = para.add_run(text)

                            # 设置字体格式
                            runs_info = data.get("runs", [])
                            if runs_info and len(runs_info) > 0:
                                run_info = runs_info[0]
                                if run_info.get("bold"):
                                    run.bold = True
                                if run_info.get("italic"):
                                    run.italic = True
                                if run_info.get("underline"):
                                    run.underline = True
                                if run_info.get("font_size"):
                                    run.font.size = run_info["font_size"]
                                if run_info.get("font_name"):
                                    run.font.name = run_info["font_name"]
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), run_info["font_name"])

                    except Exception as e:
                        print(f"处理段落失败: {e}")

                    content_idx += 1

            # 处理表格
            elif child.tag == _qn('w:tbl'):
                if table_idx < len(doc.tables):
                    table = doc.tables[table_idx]
                    table_idx += 1

                    try:
                        if isinstance(content["content_data"], str):
                            data = json.loads(content["content_data"])
                        else:
                            data = content["content_data"]

                        table_data = data.get("data", [])
                        merge_info = data.get("merge_info", [])

                        # 更新表格数据（保留单元格格式）
                        for ri, row in enumerate(table.rows):
                            if ri < len(table_data):
                                for ci, cell in enumerate(row.cells):
                                    if ci < len(table_data[ri]):
                                        # 清空单元格内容但保留格式
                                        for paragraph in cell.paragraphs[:]:
                                            for run in paragraph.runs[:]:
                                                run.text = ""
                                                run._element.getparent().remove(run._element)
                                            paragraph._element.getparent().remove(paragraph._element)
                                        
                                        # 添加新文本
                                        new_para = cell.add_paragraph()
                                        new_run = new_para.add_run(table_data[ri][ci])

                    except Exception as e:
                        print(f"处理表格失败: {e}")

                    content_idx += 1

        doc.save(output_path)
        return True


class ExcelHandler:
    """Excel文档处理工具类"""

    @staticmethod
    def parse_excel_template(file_path: str) -> Dict[str, Any]:
        """解析Excel模板，提取结构和数据"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        df = pd.read_excel(file_path, sheet_name=None)
        result = {}
        for sheet_name, sheet_df in df.items():
            # 将NaN替换为空字符串，确保JSON序列化正确
            sheet_df = sheet_df.fillna('')
            result[sheet_name] = {
                "columns": sheet_df.columns.tolist(),
                "row_count": len(sheet_df),
                "data": sheet_df.to_dict(orient="records")
            }
        return result

    @staticmethod
    def fill_word_template_with_excel(
        word_template_path: str,
        excel_data: Dict[str, Any],
        output_path: str,
        base_code: str,
        base_name: str
    ) -> str:
        """使用Excel数据填充Word模板生成申报文档"""
        if not os.path.exists(word_template_path):
            raise FileNotFoundError(f"Word模板文件不存在: {word_template_path}")

        doc = Document(word_template_path)

        # 替换段落中的占位符
        for para in doc.paragraphs:
            for key, value in excel_data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in para.text:
                    para.text = para.text.replace(placeholder, str(value))
            # 替换基地相关占位符
            para.text = para.text.replace("{{base_code}}", base_code)
            para.text = para.text.replace("{{base_name}}", base_name)

        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in excel_data.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))
                    cell.text = cell.text.replace("{{base_code}}", base_code)
                    cell.text = cell.text.replace("{{base_name}}", base_name)

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        return output_path

    @staticmethod
    def batch_generate_documents(
        word_template_path: str,
        excel_path: str,
        output_dir: str,
        base_code_col: str = "base_code",
        base_name_col: str = "base_name"
    ) -> List[Dict[str, str]]:
        """批量生成申报文档"""
        df = pd.read_excel(excel_path)
        results = []

        for _, row in df.iterrows():
            base_code = str(row.get(base_code_col, ""))
            base_name = str(row.get(base_name_col, ""))

            if not base_code:
                continue

            # 构建替换数据（排除基地编码和名称列）
            replace_data = {
                k: str(v) for k, v in row.items()
                if k not in [base_code_col, base_name_col]
            }

            output_path = os.path.join(
                output_dir, base_code,
                f"申报文档_{base_code}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
            )

            try:
                ExcelHandler.fill_word_template_with_excel(
                    word_template_path, replace_data, output_path,
                    base_code, base_name
                )
                results.append({
                    "base_code": base_code,
                    "base_name": base_name,
                    "file_path": output_path,
                    "status": "success"
                })
            except Exception as e:
                results.append({
                    "base_code": base_code,
                    "base_name": base_name,
                    "file_path": "",
                    "status": f"failed: {str(e)}"
                })

        return results
